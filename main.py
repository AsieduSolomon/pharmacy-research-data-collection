"""
Drug Storage Conditions Survey
Sunyani Technical University — Department of Pharmacy
Supabase · Rich Charts · PDF & Word Report Export
"""

import streamlit as st
import pandas as pd
import plotly.express as px
import plotly.graph_objects as go
import re
from datetime import datetime
from io import BytesIO
from collections import Counter

from supabase import create_client, Client
SUPABASE_URL = st.secrets["SUPABASE_URL"]
SUPABASE_KEY = st.secrets["SUPABASE_KEY"]

@st.cache_resource
def get_supabase() -> Client:
    return create_client(SUPABASE_URL, SUPABASE_KEY)

supabase: Client = get_supabase()

st.set_page_config(page_title="Drug Storage Survey | STU", page_icon="💊",
                   layout="wide", initial_sidebar_state="expanded")

TEAL="#028090"; GREEN="#02C39A"; NAVY="#065A82"; CORAL="#F96167"
GOLD="#F9C74F"; MINT="#90E0EF"; SLATE="#1C7293"; AMBER="#F4A261"; PURPLE="#7B2D8B"
PAL=[TEAL,GREEN,CORAL,GOLD,NAVY,AMBER,MINT,SLATE,PURPLE]

st.markdown("""
<style>
@import url('https://fonts.googleapis.com/css2?family=Inter:wght@400;600;700;800&display=swap');
html,body,[class*="css"]{font-family:'Inter',sans-serif;}
.main-header{background:linear-gradient(135deg,#028090 0%,#065A82 55%,#21295C 100%);
  padding:32px 40px;border-radius:16px;margin-bottom:28px;color:white;text-align:center;
  box-shadow:0 8px 32px rgba(2,128,144,.30);}
.main-header h1{font-size:1.75em;margin:0;font-weight:800;letter-spacing:-.5px;}
.main-header p{font-size:.92em;margin:6px 0 0;opacity:.88;}
.section-card{background:linear-gradient(90deg,#f0fafc,#f8fdfe);border-left:6px solid #028090;
  border-radius:10px;padding:14px 20px;margin:22px 0 10px;
  box-shadow:0 2px 8px rgba(2,128,144,.08);}
.section-title{color:#065A82;font-size:1.0em;font-weight:700;text-transform:uppercase;letter-spacing:.8px;}
.kpi-card{background:white;border-radius:14px;padding:22px 14px;text-align:center;
  box-shadow:0 4px 18px rgba(0,0,0,.07);border-top:5px solid #028090;margin-bottom:10px;}
.kpi-value{font-size:2.2em;font-weight:800;color:#028090;}
.kpi-label{font-size:.80em;color:#666;margin-top:5px;font-weight:600;}
.kpi-green{border-top-color:#02C39A;} .kpi-green .kpi-value{color:#02C39A;}
.kpi-coral{border-top-color:#F96167;} .kpi-coral .kpi-value{color:#F96167;}
.kpi-gold{border-top-color:#F9C74F;}  .kpi-gold .kpi-value{color:#c49a00;}
.kpi-navy{border-top-color:#065A82;}  .kpi-navy .kpi-value{color:#065A82;}
.success-box{background:linear-gradient(135deg,#e8faf7,#d0f5ee);border:2px solid #02C39A;
  border-radius:14px;padding:36px;text-align:center;margin:24px 0;}
.info-banner{background:linear-gradient(90deg,#028090,#00A896);color:white;
  padding:13px 22px;border-radius:10px;margin-bottom:18px;font-size:.92em;}
.admin-banner{background:linear-gradient(135deg,#21295C,#065A82);color:white;
  padding:20px 26px;border-radius:12px;margin-bottom:22px;
  border-left:6px solid #F9C74F;}
.chart-insight{background:linear-gradient(90deg,#f0fafc,#e8f7fb);
  border-left:4px solid #028090;border-radius:8px;
  padding:10px 16px;font-size:.85em;color:#065A82;margin-top:6px;font-weight:500;}
.section-divider{border:none;border-top:2px solid #e0f0f4;margin:32px 0;}
footer{visibility:hidden;}
</style>
""", unsafe_allow_html=True)

# ── DB helpers ───────────────────────────────────────────────────
def insert_response(data):
    try:
        supabase.table("responses").insert(data).execute()
        return True
    except Exception as e:
        st.error(f"Database error: {e}")
        return False

def fetch_all():
    try:
        res = supabase.table("responses").select("*").execute()
        return pd.DataFrame(res.data) if res.data else pd.DataFrame()
    except Exception as e:
        st.error(f"Failed to fetch data: {e}")
        return pd.DataFrame()

def delete_response(rid):
    supabase.table("responses").delete().eq("id", rid).execute()

def response_count():
    try:
        res = supabase.table("responses").select("id", count="exact").execute()
        return res.count or 0
    except:
        return 0

def prep_df(df):
    bool_cols=["has_ac","has_refrigerator","has_thermometer","has_hygrometer",
               "has_proper_shelving","has_ventilation","has_direct_sunlight",
               "has_written_sop","gsp_training_received","fda_inspected",
               "observed_degradation","observed_color_change","potency_complaints","returned_stock_quality"]
    for c in bool_cols:
        if c in df.columns:
            df[c]=pd.to_numeric(df[c],errors="coerce").fillna(0).astype(int)
    return df

# ── Chart helpers ────────────────────────────────────────────────
LAYOUT_BASE=dict(font=dict(family="Inter,Arial,sans-serif",size=12,color="#333"),
  plot_bgcolor="white",paper_bgcolor="white",margin=dict(l=20,r=20,t=70,b=20),
  hoverlabel=dict(bgcolor="white",font_size=13,font_family="Inter"))

def apply_layout(fig,title,subtitle="",height=400,**kw):
    ttl=f"<b>{title}</b>"+(f"<br><sup style='color:#666'>{subtitle}</sup>" if subtitle else "")
    fig.update_layout(**LAYOUT_BASE,
      title=dict(text=ttl,font=dict(size=15,color="#065A82"),x=0.01,xanchor="left"),
      height=height,**kw)
    fig.update_xaxes(showgrid=True,gridcolor="#f0f0f0",zeroline=False)
    fig.update_yaxes(showgrid=True,gridcolor="#f0f0f0",zeroline=False)
    return fig

def insight(text):
    st.markdown(f'<div class="chart-insight">💡 {text}</div>',unsafe_allow_html=True)

# ── Auth ─────────────────────────────────────────────────────────
ADMIN_PASSWORD="kwame"

def require_admin():
    if st.session_state.get("admin_authenticated"):
        return True
    st.markdown('<div class="admin-banner">🔐 <b>Admin Panel</b> — Restricted Access<br>'
                '<small>Enter the admin password to continue.</small></div>',unsafe_allow_html=True)
    pwd=st.text_input("Admin Password",type="password",key="admin_pwd_input")
    if st.button("🔓 Unlock Admin Panel",use_container_width=True):
        if pwd==ADMIN_PASSWORD:
            st.session_state["admin_authenticated"]=True
            st.rerun()
        else:
            st.error("❌ Incorrect password.")
    return False

def is_valid_email(e):
    return bool(re.match(r"^[\w\.\+\-]+@[\w\-]+\.[a-zA-Z]{2,}$",e))

def render_header():
    st.markdown("""
    <div class="main-header">
      <h1>💊 Drug Storage Conditions Survey</h1>
      <p><strong>Sunyani Technical University</strong> &nbsp;|&nbsp; Department of Pharmacy</p>
      <p style="font-size:.78em;opacity:.75;margin-top:7px;">
        Evaluation of Storage Conditions of Drugs in Community Pharmacies and Their
        Influence on Drug Potency and Shelf Life in the Sunyani Municipality, Ghana — 2026
      </p>
    </div>""",unsafe_allow_html=True)

# ════════════════════════════════════════════════════════
# PAGE: SURVEY
# ════════════════════════════════════════════════════════
def page_survey():
    st.markdown('<div class="info-banner">📋 &nbsp;Please answer all questions honestly. '
                'Responses are <strong>confidential</strong> and for academic research only. '
                'Fields marked <strong>*</strong> are required.</div>',unsafe_allow_html=True)
    with st.form("survey_form",clear_on_submit=True):
        st.markdown('<div class="section-card"><div class="section-title">🏪 Section A — Pharmacy & Pharmacist Profile</div></div>',unsafe_allow_html=True)
        c1,c2=st.columns(2)
        with c1:
            pharmacy_name   =st.text_input("Name of Pharmacy *",placeholder="e.g. Sunshine Pharmacy")
            pharmacist_name =st.text_input("Name of Pharmacist / Respondent *",placeholder="Full name")
            reg_number      =st.text_input("Pharmacy Council Registration Number *",placeholder="e.g. PC/2019/001")
            pharmacy_type   =st.selectbox("Type of Pharmacy *",["-- Select --","Retail Community Pharmacy","Hospital Pharmacy","Licensed Chemical Shop","Wholesale Pharmacy"])
        with c2:
            pharmacist_email=st.text_input("Email Address (optional)",placeholder="name@example.com")
            years_experience=st.number_input("Years of Professional Experience *",min_value=0,max_value=60,step=1)
            num_staff       =st.number_input("Total Number of Staff *",min_value=1,max_value=200,step=1)
            location_type   =st.selectbox("Pharmacy Location *",["-- Select --","Urban (City Centre)","Peri-urban","Rural"])

        st.markdown('<div class="section-card"><div class="section-title">🌡️ Section B — Physical Storage Conditions & Facilities</div></div>',unsafe_allow_html=True)
        st.markdown("**Does your pharmacy have the following equipment?**")
        c1,c2,c3=st.columns(3)
        with c1:
            has_ac          =st.radio("Air Conditioning Unit *",["Yes","No"],horizontal=True)
            has_refrigerator=st.radio("Functional Refrigerator *",["Yes","No"],horizontal=True)
        with c2:
            has_thermometer =st.radio("Room Thermometer *",["Yes","No"],horizontal=True)
            has_hygrometer  =st.radio("Hygrometer (Humidity Meter) *",["Yes","No"],horizontal=True)
        with c3:
            has_proper_shelving=st.radio("Proper Drug Shelving *",["Yes","No"],horizontal=True)
            has_ventilation    =st.radio("Adequate Ventilation *",["Yes","No"],horizontal=True)
        c1,c2=st.columns(2)
        with c1: has_direct_sunlight=st.radio("Drugs exposed to direct sunlight? *",["Yes","No"],horizontal=True)
        with c2: storage_area_size=st.selectbox("Size of Main Drug Storage Area *",["-- Select --","Very small (< 10 m²)","Small (10–25 m²)","Medium (25–50 m²)","Large (> 50 m²)"])

        st.markdown('<div class="section-card"><div class="section-title">📊 Section C — Temperature & Humidity Monitoring</div></div>',unsafe_allow_html=True)
        c1,c2=st.columns(2)
        with c1:
            temp_monitoring_freq=st.selectbox("How often is room temperature recorded? *",["-- Select --","Multiple times daily","Once daily","Weekly","Monthly","Never monitored"])
            usual_temp_range    =st.selectbox("Usual temperature range in storage area *",["-- Select --","Below 25°C  (WHO Optimal)","25°C – 30°C  (Acceptable)","30°C – 35°C  (Risk Zone)","Above 35°C  (High Risk)","Not monitored / Unknown"])
        with c2:
            usual_humidity_range =st.selectbox("Usual relative humidity (RH%) *",["-- Select --","Below 45% RH  (Low)","45% – 65% RH  (Optimal)","65% – 75% RH  (Elevated Risk)","Above 75% RH  (High Risk)","Not monitored / Unknown"])
            temp_excursion_action=st.selectbox("Action taken when temperature exceeds limits *",["-- Select --","Move drugs to refrigerator / cooler area","Increase ventilation or switch on AC","Document and notify supervisor","No specific action taken","No system in place to detect excursions"])
        has_written_sop=st.radio("Does your pharmacy have a written SOP for drug storage? *",["Yes","No"],horizontal=True)

        st.markdown('<div class="section-card"><div class="section-title">✅ Section D — Drug Handling & WHO GSP Compliance</div></div>',unsafe_allow_html=True)
        c1,c2=st.columns(2)
        with c1:
            gsp_training_received=st.radio("Have you received WHO GSP training? *",["Yes","No"],horizontal=True)
            last_training_year   =st.selectbox("If yes — most recent training?",["N/A — No training received","2024 – 2025","2021 – 2023","2018 – 2020","Before 2018"])
            practises_fifo       =st.radio("Do you practise FIFO stock rotation? *",["Yes","No","Sometimes"],horizontal=True)
        with c2:
            checks_expiry          =st.radio("Do you regularly inspect for expiry dates? *",["Yes","No","Sometimes"],horizontal=True)
            segregates_thermolabile=st.radio("Temperature-sensitive drugs stored separately under cold conditions? *",["Yes","No","Not Applicable"],horizontal=True)
            fda_inspected          =st.radio("Inspected by FDA Ghana in the last 2 years? *",["Yes","No"],horizontal=True)
        self_compliance_rating=st.select_slider("Rate your pharmacy's overall WHO GSP compliance *",options=["Very Poor","Poor","Fair","Good","Excellent"],value="Fair")

        st.markdown('<div class="section-card"><div class="section-title">⚠️ Section E — Observed Drug Quality & Potency Issues</div></div>',unsafe_allow_html=True)
        c1,c2=st.columns(2)
        with c1:
            observed_degradation  =st.radio("Observed physical signs of drug degradation? *",["Yes","No"],horizontal=True)
            degradation_drug_types=st.multiselect("If yes — which drug types?",["Antibiotics (tablets/capsules)","Syrups / Oral liquids","Antimalarials","Antihypertensives","Injectables / Vaccines","Suppositories / Pessaries","Topical creams / ointments","Other"])
            observed_color_change =st.radio("Noticed unusual colour change in any drug? *",["Yes","No"],horizontal=True)
        with c2:
            potency_complaints    =st.radio("Patients complained medicine was not working? *",["Yes","No"],horizontal=True)
            returned_stock_quality=st.radio("Returned or discarded stock due to quality deterioration? *",["Yes","No"],horizontal=True)
            num_quality_incidents =st.selectbox("Drug quality incidents in the last 12 months? *",["-- Select --","None","1 – 3","4 – 6","7 – 10","More than 10"])

        st.markdown('<div class="section-card"><div class="section-title">💡 Section F — Challenges & Recommendations</div></div>',unsafe_allow_html=True)
        biggest_challenge=st.selectbox("Biggest challenge to maintaining proper drug storage? *",["-- Select --","Frequent power outages / erratic electricity supply","High cost of air conditioning or refrigeration","Small or inadequate storage space","Lack of temperature/humidity monitoring equipment","Limited training and awareness on GSP","Insufficient regulatory inspection and enforcement","Financial constraints to upgrade facilities"])
        support_needed=st.selectbox("Support that would most improve drug storage? *",["-- Select --","Subsidised air conditioning and refrigeration equipment","Regular GSP training workshops","Affordable data loggers or thermometers","More frequent FDA inspections and guidance","Financial support / grants for facility upgrades","National policy changes and stricter enforcement"])
        policy_recommendation=st.text_area("Policy recommendations? (optional)",placeholder="Share your suggestions...",height=90)
        additional_comments  =st.text_area("Additional comments? (optional)",placeholder="Anything else you would like to share...",height=80)

        st.markdown("---")
        st.markdown("**⚠️ Please review all answers before submitting.**")
        submitted=st.form_submit_button("✅  Submit Survey Response",use_container_width=True)

        if submitted:
            errors=[]
            if not pharmacy_name.strip():   errors.append("Pharmacy name is required.")
            if not pharmacist_name.strip(): errors.append("Pharmacist name is required.")
            if not reg_number.strip():      errors.append("Registration number is required.")
            if pharmacist_email.strip() and not is_valid_email(pharmacist_email):
                errors.append("Please enter a valid email address.")
            for lbl,val in {"Type of pharmacy":pharmacy_type,"Pharmacy location":location_type,
                "Storage area size":storage_area_size,"Temperature monitoring frequency":temp_monitoring_freq,
                "Usual temperature range":usual_temp_range,"Usual humidity range":usual_humidity_range,
                "Action on temperature excursion":temp_excursion_action,
                "Number of quality incidents":num_quality_incidents,
                "Biggest challenge":biggest_challenge,"Support needed":support_needed}.items():
                if val.startswith("-- Select"): errors.append(f"'{lbl}' is required.")
            if errors:
                st.error("**Please fix the following before submitting:**")
                for e in errors: st.markdown(f"- ❌ {e}")
            else:
                data={
                    "submitted_at":datetime.now().strftime("%Y-%m-%d %H:%M:%S"),
                    "pharmacy_name":pharmacy_name.strip(),"pharmacist_name":pharmacist_name.strip(),
                    "pharmacist_email":pharmacist_email.strip(),"reg_number":reg_number.strip(),
                    "years_experience":int(years_experience),"pharmacy_type":pharmacy_type,
                    "num_staff":int(num_staff),"location_type":location_type,
                    "has_ac":1 if has_ac=="Yes" else 0,"has_refrigerator":1 if has_refrigerator=="Yes" else 0,
                    "has_thermometer":1 if has_thermometer=="Yes" else 0,"has_hygrometer":1 if has_hygrometer=="Yes" else 0,
                    "has_proper_shelving":1 if has_proper_shelving=="Yes" else 0,"has_ventilation":1 if has_ventilation=="Yes" else 0,
                    "has_direct_sunlight":1 if has_direct_sunlight=="Yes" else 0,"storage_area_size":storage_area_size,
                    "temp_monitoring_freq":temp_monitoring_freq,"usual_temp_range":usual_temp_range,
                    "usual_humidity_range":usual_humidity_range,"temp_excursion_action":temp_excursion_action,
                    "has_written_sop":1 if has_written_sop=="Yes" else 0,
                    "gsp_training_received":1 if gsp_training_received=="Yes" else 0,
                    "last_training_year":last_training_year,"practises_fifo":practises_fifo,
                    "checks_expiry":checks_expiry,"segregates_thermolabile":segregates_thermolabile,
                    "fda_inspected":1 if fda_inspected=="Yes" else 0,"self_compliance_rating":self_compliance_rating,
                    "observed_degradation":1 if observed_degradation=="Yes" else 0,
                    "degradation_drug_types":", ".join(degradation_drug_types),
                    "observed_color_change":1 if observed_color_change=="Yes" else 0,
                    "potency_complaints":1 if potency_complaints=="Yes" else 0,
                    "returned_stock_quality":1 if returned_stock_quality=="Yes" else 0,
                    "num_quality_incidents":num_quality_incidents,"biggest_challenge":biggest_challenge,
                    "support_needed":support_needed,"policy_recommendation":policy_recommendation.strip(),
                    "additional_comments":additional_comments.strip(),
                }
                if insert_response(data):
                    st.markdown("""
                    <div class="success-box">
                      <div style="font-size:3em;">✅</div>
                      <h2 style="color:#028090;margin:10px 0;">Response Submitted Successfully!</h2>
                      <p style="color:#065A82;font-size:1.05em;">
                        Thank you for completing this survey.<br>
                        Your contribution is vital to improving drug storage practices
                        and patient safety in the Sunyani Municipality.
                      </p>
                      <p style="font-size:.82em;color:#888;margin-top:14px;">
                        Sunyani Technical University — Department of Pharmacy, 2026
                      </p>
                    </div>""",unsafe_allow_html=True)

# ════════════════════════════════════════════════════════
# PAGE: DASHBOARD
# ════════════════════════════════════════════════════════
def show(fig, ins=""):
    """Render a plotly figure + optional insight box."""
    st.plotly_chart(fig, use_container_width=True, config={"displayModeBar": False})
    if ins:
        st.markdown(f'<div class="chart-insight">💡 {ins}</div>', unsafe_allow_html=True)

def page_dashboard():
    df=fetch_all()
    if df.empty:
        st.warning("📭 No survey responses yet. Share the survey link with pharmacists.")
        return
    df=prep_df(df); n=len(df)

    # ── KPI Row ──────────────────────────────────────────
    st.markdown("### 📈 Key Performance Indicators")
    c1,c2,c3,c4,c5,c6=st.columns(6)
    kpis=[
        (c1,str(n),"Total Responses",""),
        (c2,f"{df['has_ac'].mean()*100:.0f}%","Have Air Conditioning","kpi-coral" if df['has_ac'].mean()<.7 else "kpi-green"),
        (c3,f"{df['has_thermometer'].mean()*100:.0f}%","Have Thermometer","kpi-coral" if df['has_thermometer'].mean()<.7 else "kpi-green"),
        (c4,f"{df['gsp_training_received'].mean()*100:.0f}%","GSP Trained","kpi-gold"),
        (c5,f"{df['observed_degradation'].mean()*100:.0f}%","Drug Degradation","kpi-coral"),
        (c6,f"{df['has_written_sop'].mean()*100:.0f}%","Have Written SOP","kpi-navy"),
    ]
    for col,val,label,cls in kpis:
        with col:
            st.markdown(f'<div class="kpi-card {cls}"><div class="kpi-value">{val}</div>'
                        f'<div class="kpi-label">{label}</div></div>',unsafe_allow_html=True)

    st.markdown('<hr class="section-divider">',unsafe_allow_html=True)

    # ══ SECTION 1: Infrastructure ══════════════════════
    st.markdown("## 🏗️ Section 1 — Storage Infrastructure")
    c1,c2=st.columns(2)
    with c1:
        labels=["Air Conditioning","Refrigerator","Thermometer","Hygrometer","Proper Shelving","Ventilation"]
        keys  =["has_ac","has_refrigerator","has_thermometer","has_hygrometer","has_proper_shelving","has_ventilation"]
        vals  =[round(df[k].mean()*100,1) for k in keys]
        colors=[GREEN if v>=70 else CORAL for v in vals]
        fig=go.Figure(go.Bar(y=labels,x=vals,orientation='h',
            marker=dict(color=colors,line=dict(width=0)),
            text=[f"  {v}%" for v in vals],textposition='outside',
            textfont=dict(size=12,color="#333"),
            hovertemplate='<b>%{y}</b><br>%{x:.1f}% of pharmacies<extra></extra>'))
        fig.add_vline(x=70,line_dash="dot",line_color=NAVY,line_width=2.5,
            annotation_text=" 70% Benchmark",annotation_position="top right",
            annotation_font=dict(size=11,color=NAVY))
        apply_layout(fig,"Storage Equipment Availability","Green ≥70% adequate | Red <70% critical gap",
            height=380,xaxis=dict(range=[0,125],title="% of Pharmacies"),yaxis_title="")
        show(fig,"Red bars show critical equipment gaps — each is a direct risk to drug potency and patient safety.")
    with c2:
        sz=df["storage_area_size"].value_counts().reset_index(); sz.columns=["Size","Count"]
        fig2=go.Figure(go.Pie(labels=sz["Size"],values=sz["Count"],hole=.45,
            marker=dict(colors=PAL,line=dict(color="white",width=3)),
            textinfo='percent+label',textfont_size=12,pull=[.03]*len(sz),
            hovertemplate='<b>%{label}</b><br>%{value} pharmacies (%{percent})<extra></extra>'))
        apply_layout(fig2,"Pharmacy Storage Area Sizes","Distribution of drug storage space",
            height=380,showlegend=False)
        fig2.update_xaxes(visible=False); fig2.update_yaxes(visible=False)
        show(fig2,"Small storage areas increase risk of improper stacking, poor ventilation and cross-contamination.")

    c1,c2=st.columns(2)
    with c1:
        pt=df["pharmacy_type"].value_counts().reset_index(); pt.columns=["Type","Count"]
        pct=(pt["Count"]/n*100).round(1)
        fig3=go.Figure(go.Bar(x=pt["Type"],y=pt["Count"],
            marker=dict(color=PAL[:len(pt)],line=dict(width=0)),
            text=[f"{c}<br>({p}%)" for c,p in zip(pt["Count"],pct)],textposition='outside',
            hovertemplate='<b>%{x}</b><br>%{y} pharmacies<extra></extra>'))
        apply_layout(fig3,"Pharmacy Types Surveyed","Breakdown by licence category",
            height=360,xaxis_title="Pharmacy Type",yaxis_title="Count",xaxis_tickangle=-15)
        show(fig3,"Retail community pharmacies form the majority — reflecting their dominance in Sunyani Municipality.")
    with c2:
        lt=df["location_type"].value_counts().reset_index(); lt.columns=["Location","Count"]
        pct_l=(lt["Count"]/n*100).round(1)
        cl=[TEAL if "Urban" in l else AMBER if "Peri" in l else CORAL for l in lt["Location"]]
        fig4=go.Figure(go.Bar(x=lt["Location"],y=lt["Count"],
            marker=dict(color=cl,line=dict(width=0)),
            text=[f"{c}<br>({p}%)" for c,p in zip(lt["Count"],pct_l)],textposition='outside',
            hovertemplate='<b>%{x}</b><br>%{y} pharmacies<extra></extra>'))
        apply_layout(fig4,"Pharmacy Location Distribution","Urban vs peri-urban vs rural representation",
            height=360,xaxis_title="Location Type",yaxis_title="Count")
        show(fig4,"Rural pharmacies face greater infrastructure challenges — even small rural samples matter for equity analysis.")

    st.markdown('<hr class="section-divider">',unsafe_allow_html=True)

    # ══ SECTION 2: Temperature & Humidity ══════════════
    st.markdown("## 🌡️ Section 2 — Temperature & Humidity Conditions")
    c1,c2=st.columns(2)
    with c1:
        tmap={"Below 25°C  (WHO Optimal)":"<25°C ✅ Optimal","25°C – 30°C  (Acceptable)":"25-30°C ⚠️ Acceptable",
              "30°C – 35°C  (Risk Zone)":"30-35°C 🔶 Risk Zone","Above 35°C  (High Risk)":">35°C 🔴 High Risk","Not monitored / Unknown":"Unknown"}
        df["tlbl"]=df["usual_temp_range"].map(tmap).fillna(df["usual_temp_range"])
        tc=df["tlbl"].value_counts().reset_index(); tc.columns=["Range","Count"]
        tcol=[GREEN if "Optimal" in r else TEAL if "Acceptable" in r else GOLD if "Risk Zone" in r else CORAL if "High Risk" in r else SLATE for r in tc["Range"]]
        fig5=go.Figure(go.Pie(labels=tc["Range"],values=tc["Count"],hole=.50,
            marker=dict(colors=tcol,line=dict(color="white",width=3)),
            textinfo='percent+label',textfont_size=12,pull=[.04]*len(tc),
            hovertemplate='<b>%{label}</b><br>%{value} pharmacies (%{percent})<extra></extra>'))
        apply_layout(fig5,"Usual Temperature in Storage Areas","WHO recommends ≤25-30°C for most medicines",
            height=400,showlegend=False)
        fig5.update_xaxes(visible=False); fig5.update_yaxes(visible=False)
        show(fig5,"Orange and red segments = pharmacies above WHO safe limits — drugs face accelerated chemical degradation.")
    with c2:
        hmap={"Below 45% RH  (Low)":"<45% RH (Low)","45% – 65% RH  (Optimal)":"45-65% ✅ Optimal",
              "65% – 75% RH  (Elevated Risk)":"65-75% ⚠️ Elevated","Above 75% RH  (High Risk)":">75% 🔴 High Risk","Not monitored / Unknown":"Unknown"}
        df["hlbl"]=df["usual_humidity_range"].map(hmap).fillna(df["usual_humidity_range"])
        hc=df["hlbl"].value_counts().reset_index(); hc.columns=["Humidity","Count"]
        hcol=[GREEN if "Optimal" in r else GOLD if "Elevated" in r else CORAL if "High" in r else SLATE for r in hc["Humidity"]]
        fig6=go.Figure(go.Pie(labels=hc["Humidity"],values=hc["Count"],hole=0,
            marker=dict(colors=hcol,line=dict(color="white",width=3)),
            textinfo='percent+label',textfont_size=12,
            hovertemplate='<b>%{label}</b><br>%{value} pharmacies (%{percent})<extra></extra>'))
        apply_layout(fig6,"Relative Humidity in Storage Areas","Optimal range: 45-65% RH",
            height=400,showlegend=False)
        fig6.update_xaxes(visible=False); fig6.update_yaxes(visible=False)
        show(fig6,"High humidity (>65% RH) accelerates moisture-induced degradation in antibiotics and oral tablets.")

    c1,c2=st.columns(2)
    with c1:
        fmap={"Multiple times daily":"Multiple/Day","Once daily":"Once Daily","Weekly":"Weekly","Monthly":"Monthly","Never monitored":"Never"}
        ford=["Multiple/Day","Once Daily","Weekly","Monthly","Never"]
        fcol=[GREEN,GREEN,GOLD,AMBER,CORAL]
        df["fshort"]=df["temp_monitoring_freq"].map(fmap).fillna(df["temp_monitoring_freq"])
        fc=df["fshort"].value_counts().reindex(ford,fill_value=0).reset_index(); fc.columns=["Frequency","Count"]
        fig7=go.Figure(go.Bar(x=fc["Frequency"],y=fc["Count"],
            marker=dict(color=fcol,line=dict(width=0)),
            text=fc["Count"],textposition='outside',textfont_size=13,
            hovertemplate='<b>%{x}</b><br>%{y} pharmacies<extra></extra>'))
        apply_layout(fig7,"Temperature Monitoring Frequency","WHO GSP requires minimum once-daily recording",
            height=360,xaxis_title="Monitoring Frequency",yaxis_title="Pharmacies")
        show(fig7,"'Weekly', 'Monthly' or 'Never' monitors cannot detect dangerous temperature excursions in time.")
    with c2:
        if "submitted_at" in df.columns:
            df["sdt"]=pd.to_datetime(df["submitted_at"],errors="coerce")
            daily=df.dropna(subset=["sdt"]).groupby(df["sdt"].dt.date).size().reset_index()
            daily.columns=["Date","Count"]; daily=daily.sort_values("Date")
            daily["Cumulative"]=daily["Count"].cumsum()
            fig8=go.Figure()
            fig8.add_trace(go.Bar(x=daily["Date"],y=daily["Count"],name="Daily",
                marker_color=MINT,opacity=.7,hovertemplate='<b>%{x}</b><br>%{y} new responses<extra></extra>'))
            fig8.add_trace(go.Scatter(x=daily["Date"],y=daily["Cumulative"],name="Cumulative",
                mode='lines+markers',line=dict(color=TEAL,width=3),
                marker=dict(size=9,color=GREEN,line=dict(width=2,color=NAVY)),
                hovertemplate='<b>%{x}</b><br>Total: %{y} responses<extra></extra>'))
            apply_layout(fig8,"Survey Submission Trend Over Time","Daily (bars) and cumulative total (line)",
                height=360,xaxis_title="Date",yaxis_title="Responses",
                legend=dict(orientation='h',y=-.25))
            show(fig8,"This trend tracks data collection progress — useful for reporting fieldwork timeline in thesis methodology.")

    st.markdown('<hr class="section-divider">',unsafe_allow_html=True)

    # ══ SECTION 3: GSP Compliance ══════════════════════
    st.markdown("## ✅ Section 3 — WHO Good Storage Practice (GSP) Compliance")
    c1,c2=st.columns(2)
    with c1:
        clbl=["GSP Training","Written SOP","FIFO","Expiry Checks","Cold Chain","FDA Inspected"]
        yp=[df["gsp_training_received"].mean()*100,df["has_written_sop"].mean()*100,
            (df["practises_fifo"]=="Yes").mean()*100,(df["checks_expiry"]=="Yes").mean()*100,
            (df["segregates_thermolabile"]=="Yes").mean()*100,df["fda_inspected"].mean()*100]
        np_=[100-v for v in yp]
        fig9=go.Figure()
        fig9.add_trace(go.Bar(name="✅ Compliant",x=clbl,y=yp,marker=dict(color=GREEN,line=dict(width=0)),
            text=[f"{v:.0f}%" for v in yp],textposition='inside',insidetextanchor='middle',
            textfont=dict(color="white",size=12)))
        fig9.add_trace(go.Bar(name="❌ Non-compliant",x=clbl,y=np_,marker=dict(color=CORAL,line=dict(width=0)),
            text=[f"{v:.0f}%" for v in np_],textposition='inside',insidetextanchor='middle',
            textfont=dict(color="white",size=12)))
        apply_layout(fig9,"WHO GSP Compliance Across All Indicators","Green=compliant | Red=non-compliant",
            height=420,barmode='stack',xaxis_tickangle=-15,
            yaxis=dict(title="% of Pharmacies",range=[0,100]),
            legend=dict(orientation='h',y=-.28))
        show(fig9,"Large red segments show which GSP practices are most neglected — priority areas for training and enforcement.")
    with c2:
        cats=["Degradation\nObserved","Colour\nChanges","Potency\nComplaints","Returned\nStock","No Written\nSOP","No\nThermometer"]
        rv=[df["observed_degradation"].mean()*100,df["observed_color_change"].mean()*100,
            df["potency_complaints"].mean()*100,df["returned_stock_quality"].mean()*100,
            (1-df["has_written_sop"].mean())*100,(1-df["has_thermometer"].mean())*100]
        fig10=go.Figure(go.Scatterpolar(r=rv+[rv[0]],theta=cats+[cats[0]],fill='toself',
            fillcolor="rgba(249,97,103,.18)",line=dict(color=CORAL,width=3),
            marker=dict(size=8,color=CORAL),
            hovertemplate='<b>%{theta}</b><br>%{r:.1f}% of pharmacies<extra></extra>'))
        fig10.update_layout(**LAYOUT_BASE,
            polar=dict(radialaxis=dict(visible=True,range=[0,100],ticksuffix="%",
                tickfont=dict(size=10),gridcolor="#eee"),
                angularaxis=dict(tickfont=dict(size=11))),
            title=dict(text="<b>Drug Quality Risk Profile</b><br>"
                "<sup style='color:#666'>Larger shaded area = broader systemic risk</sup>",
                font=dict(size=15,color="#065A82"),x=.01),
            height=420,showlegend=False)
        show(fig10,"A large shaded area means risk is widespread across multiple dimensions — not just one isolated problem.")

    c1,c2=st.columns(2)
    with c1:
        order=["Very Poor","Poor","Fair","Good","Excellent"]
        rc=df["self_compliance_rating"].value_counts().reindex(order,fill_value=0).reset_index(); rc.columns=["Rating","Count"]
        pct_rc=(rc["Count"]/n*100).round(1)
        cmap={"Very Poor":CORAL,"Poor":AMBER,"Fair":GOLD,"Good":TEAL,"Excellent":GREEN}
        fig11=go.Figure(go.Bar(x=rc["Rating"],y=rc["Count"],
            marker=dict(color=[cmap[r] for r in rc["Rating"]],line=dict(width=0)),
            text=[f"{c} ({p}%)" for c,p in zip(rc["Count"],pct_rc)],
            textposition='outside',textfont_size=12,
            hovertemplate='<b>%{x}</b><br>%{y} pharmacies<extra></extra>'))
        apply_layout(fig11,"Pharmacists' Self-Rated GSP Compliance","Self-perception vs objective compliance data",
            height=380,xaxis_title="Self-Rating",yaxis_title="Count")
        show(fig11,"Compare this self-perception with the objective stacked bar — a large gap indicates overconfidence and training need.")
    with c2:
        rnum={"Very Poor":1,"Poor":2,"Fair":3,"Good":4,"Excellent":5}
        df["cnum"]=df["self_compliance_rating"].map(rnum)
        sdf=df.dropna(subset=["years_experience","cnum"])
        fig12=px.scatter(sdf,x="years_experience",y="cnum",color="pharmacy_type",
            color_discrete_sequence=PAL,
            labels={"years_experience":"Years of Experience","cnum":"Compliance Rating","pharmacy_type":"Type"},
            hover_data={"pharmacy_name":True,"self_compliance_rating":True,"years_experience":True},
            trendline="ols",trendline_scope="overall",trendline_color_override=NAVY)
        fig12.update_traces(marker=dict(size=10,opacity=.8,line=dict(width=1,color="white")))
        apply_layout(fig12,"Experience vs Self-Rated GSP Compliance","Each dot = one pharmacy | Navy = overall trend",
            height=380,xaxis_title="Years of Experience",
            yaxis=dict(title="Compliance (1=Very Poor, 5=Excellent)",
                tickvals=[1,2,3,4,5],ticktext=["Very Poor","Poor","Fair","Good","Excellent"]),
            legend=dict(orientation='h',y=-.28))
        show(fig12,"The trend line shows if experienced pharmacists rate themselves higher — relevant for experience-based training design.")

    st.markdown('<hr class="section-divider">',unsafe_allow_html=True)

    # ══ SECTION 4: Drug Quality ═════════════════════════
    st.markdown("## ⚠️ Section 4 — Observed Drug Quality Issues")
    c1,c2=st.columns(2)
    with c1:
        inc_ord=["None","1 – 3","4 – 6","7 – 10","More than 10"]
        ic=df["num_quality_incidents"].value_counts().reindex(inc_ord,fill_value=0).reset_index(); ic.columns=["Incidents","Count"]
        fig13=go.Figure(go.Funnel(y=ic["Incidents"],x=ic["Count"],
            textinfo="value+percent initial",textfont=dict(size=13),
            marker=dict(color=[GREEN,TEAL,GOLD,AMBER,CORAL],line=dict(color=["white"]*5,width=2)),
            connector=dict(line=dict(color="#ddd",width=1)),
            hovertemplate='<b>%{y}</b><br>%{x} pharmacies<extra></extra>'))
        apply_layout(fig13,"Drug Quality Incidents — Last 12 Months","Pharmacies reporting each incident frequency band",height=400)
        fig13.update_xaxes(visible=False); fig13.update_yaxes(visible=False)
        show(fig13,"Pharmacies reporting 4+ incidents per year signal systemic storage failures needing urgent regulatory follow-up.")
    with c2:
        qk=["observed_degradation","observed_color_change","potency_complaints","returned_stock_quality"]
        ql=["Degradation","Colour Change","Potency Complaints","Returned Stock"]
        types=df["pharmacy_type"].dropna().unique().tolist()
        fig14=go.Figure()
        for i,pt_ in enumerate(types):
            sub=df[df["pharmacy_type"]==pt_]
            v14=[round(sub[k].mean()*100,1) for k in qk]
            fig14.add_trace(go.Bar(name=pt_,x=ql,y=v14,
                marker=dict(color=PAL[i%len(PAL)],line=dict(width=0)),
                text=[f"{v}%" for v in v14],textposition='outside',
                hovertemplate=f'<b>{pt_}</b><br>%{{x}}: %{{y}}%<extra></extra>'))
        apply_layout(fig14,"Drug Quality Issues by Pharmacy Type","% of each type reporting each quality problem",
            height=400,barmode='group',xaxis_title="Quality Indicator",yaxis_title="% Reporting",
            legend=dict(orientation='h',y=-.28))
        show(fig14,"Comparing quality issue rates across pharmacy types identifies which category needs the most regulatory attention.")

    c1,c2=st.columns(2)
    with c1:
        atypes=[]
        for row in df["degradation_drug_types"].dropna():
            atypes.extend([t.strip() for t in str(row).split(",") if t.strip()])
        if atypes:
            tc_=Counter(atypes)
            tdf=pd.DataFrame(tc_.items(),columns=["Drug Type","Count"]).sort_values("Count",ascending=True)
            fig15=go.Figure(go.Bar(y=tdf["Drug Type"],x=tdf["Count"],orientation='h',
                marker=dict(color=CORAL,opacity=.85,line=dict(width=0)),
                text=tdf["Count"],textposition='outside',
                hovertemplate='<b>%{y}</b><br>%{x} pharmacies<extra></extra>'))
            apply_layout(fig15,"Most Affected Drug Categories (Degradation)","Drug types most frequently observed to degrade",
                height=380,xaxis_title="Number of Reports",yaxis_title="")
            show(fig15,"The top drug categories need priority attention in cold chain management and storage guidelines.")
        else:
            st.info("No drug type degradation data recorded yet.")
    with c2:
        ek2=["has_ac","has_refrigerator","has_thermometer","has_hygrometer","has_proper_shelving","has_ventilation"]
        el2=["AC","Fridge","Thermometer","Hygrometer","Shelving","Ventilation"]
        locs=df["location_type"].dropna().unique().tolist()
        if locs:
            hmr=[[round(df[df["location_type"]==loc][k].mean()*100,1) for k in ek2] for loc in locs]
            fig16=go.Figure(go.Heatmap(z=hmr,x=el2,y=locs,
                colorscale=[[0,CORAL],[.5,GOLD],[1,GREEN]],
                text=[[f"{v}%" for v in row] for row in hmr],
                texttemplate="%{text}",textfont=dict(size=13),
                hovertemplate='<b>%{y}</b> — <b>%{x}</b>: %{z:.1f}%<extra></extra>',
                colorbar=dict(title="% with equipment",ticksuffix="%")))
            apply_layout(fig16,"Equipment Availability by Location Type","Green=high availability | Red=critical gap",height=380)
            show(fig16,"Rural pharmacies in red cells are the most underserved — supports equity-based policy recommendations.")

    st.markdown('<hr class="section-divider">',unsafe_allow_html=True)

    # ══ SECTION 5: Barriers ════════════════════════════
    st.markdown("## 🚧 Section 5 — Challenges & Support Needs")
    c1,c2=st.columns(2)
    with c1:
        csh={"Frequent power outages / erratic electricity supply":"Power Outages",
             "High cost of air conditioning or refrigeration":"High Equipment Cost",
             "Small or inadequate storage space":"Inadequate Space",
             "Lack of temperature/humidity monitoring equipment":"No Monitoring Equipment",
             "Limited training and awareness on GSP":"Limited GSP Training",
             "Insufficient regulatory inspection and enforcement":"Weak Regulation",
             "Financial constraints to upgrade facilities":"Financial Constraints"}
        df["cshort"]=df["biggest_challenge"].map(csh).fillna(df["biggest_challenge"])
        ch=df["cshort"].value_counts().sort_values(ascending=True).reset_index(); ch.columns=["Challenge","Count"]
        pct_ch=(ch["Count"]/n*100).round(1)
        fig17=go.Figure(go.Bar(y=ch["Challenge"],x=ch["Count"],orientation='h',
            marker=dict(color=ch["Count"],colorscale=[[0,MINT],[.5,TEAL],[1,NAVY]],
                line=dict(width=0),showscale=False),
            text=[f"{c} ({p}%)" for c,p in zip(ch["Count"],pct_ch)],textposition='outside',
            hovertemplate='<b>%{y}</b><br>%{x} pharmacies<extra></extra>'))
        apply_layout(fig17,"Biggest Challenges to Proper Drug Storage","Ranked by number of pharmacies citing each barrier",
            height=400,xaxis_title="Number of Pharmacies",yaxis_title="")
        show(fig17,"The longest bar identifies the single most critical barrier — this should be the first policy intervention priority.")
    with c2:
        ssh={"Subsidised air conditioning and refrigeration equipment":"Subsidised Equipment",
             "Regular GSP training workshops":"GSP Training","Affordable data loggers or thermometers":"Affordable Monitoring",
             "More frequent FDA inspections and guidance":"More FDA Inspections",
             "Financial support / grants for facility upgrades":"Financial Grants",
             "National policy changes and stricter enforcement":"Policy Reform"}
        df["sshort"]=df["support_needed"].map(ssh).fillna(df["support_needed"])
        sp=df["sshort"].value_counts().reset_index(); sp.columns=["Support","Count"]
        fig18=go.Figure(go.Pie(labels=sp["Support"],values=sp["Count"],hole=.42,
            marker=dict(colors=PAL,line=dict(color="white",width=3)),
            textinfo='percent+label',textfont_size=12,pull=[.04]*len(sp),
            hovertemplate='<b>%{label}</b><br>%{value} pharmacies (%{percent})<extra></extra>'))
        apply_layout(fig18,"Type of Support Most Needed","What pharmacists say would most improve storage quality",
            height=400,showlegend=False)
        fig18.update_xaxes(visible=False); fig18.update_yaxes(visible=False)
        show(fig18,"The dominant segment shows where investment would have the greatest impact on drug storage quality.")

# ════════════════════════════════════════════════════════
# PAGE: RESPONSES (Admin)
# ════════════════════════════════════════════════════════
def page_responses():
    if not require_admin(): return
    df=fetch_all()
    if df.empty: st.info("📭 No responses submitted yet."); return
    st.markdown(f"### 📋 All Submitted Responses — {len(df)} total")
    sc=["id","submitted_at","pharmacy_name","pharmacist_name","pharmacy_type","location_type","self_compliance_rating","observed_degradation"]
    sc=[c for c in sc if c in df.columns]
    show=df[sc].copy()
    if "observed_degradation" in show.columns:
        show["observed_degradation"]=show["observed_degradation"].map({1:"✅ Yes",0:"❌ No"})
    show.columns=["ID","Submitted At","Pharmacy","Pharmacist","Type","Location","Compliance","Degradation?"][:len(sc)]
    st.dataframe(show,use_container_width=True,hide_index=True)
    st.markdown("---")
    st.markdown("#### 🔍 View or Delete a Response")
    rid=st.number_input("Enter Response ID",min_value=1,step=1,key="rid_input")
    c1,c2,_=st.columns([1,1,4])
    with c1:
        if st.button("👁️ View"):
            row=df[df["id"]==rid]
            if row.empty: st.error(f"No response with ID {rid}.")
            else:
                r=row.iloc[0]
                st.success(f"Details for **{r['pharmacy_name']}** — {r['submitted_at']}")
                st.json({k:str(v) for k,v in r.items() if k!="id"})
    with c2:
        if st.button("🗑️ Delete",type="secondary"):
            row=df[df["id"]==rid]
            if row.empty: st.error(f"No response with ID {rid}.")
            else:
                delete_response(int(rid)); st.success(f"Response ID {rid} deleted."); st.rerun()
    st.markdown("---")
    csv=df.to_csv(index=False).encode("utf-8")
    st.download_button("⬇️ Download All Responses as CSV",data=csv,
        file_name=f"pharmacy_survey_{datetime.now().strftime('%Y%m%d')}.csv",
        mime="text/csv",use_container_width=True)


# ════════════════════════════════════════════════════════
# PDF GENERATOR
# ════════════════════════════════════════════════════════
def generate_pdf(df,n):
    from reportlab.lib.pagesizes import A4
    from reportlab.lib import colors
    from reportlab.lib.styles import getSampleStyleSheet,ParagraphStyle
    from reportlab.lib.units import cm
    from reportlab.platypus import (SimpleDocTemplate,Paragraph,Spacer,Table,
                                    TableStyle,PageBreak,HRFlowable)
    from reportlab.lib.enums import TA_CENTER,TA_LEFT,TA_JUSTIFY
    buf=BytesIO()
    doc=SimpleDocTemplate(buf,pagesize=A4,leftMargin=2*cm,rightMargin=2*cm,topMargin=2.5*cm,bottomMargin=2*cm)
    C_TEAL=colors.HexColor("#028090"); C_NAVY=colors.HexColor("#065A82")
    C_GREEN=colors.HexColor("#02C39A"); C_CORAL=colors.HexColor("#F96167")
    C_LITE=colors.HexColor("#e8f7fb"); C_GREY=colors.HexColor("#f5f5f5"); C_DARK=colors.HexColor("#222222")
    sty={
        "ct":ParagraphStyle("ct",fontSize=22,textColor=colors.white,alignment=TA_CENTER,fontName="Helvetica-Bold",leading=28),
        "cs":ParagraphStyle("cs",fontSize=12,textColor=colors.white,alignment=TA_CENTER,fontName="Helvetica",leading=18),
        "h1":ParagraphStyle("h1",fontSize=15,textColor=C_NAVY,fontName="Helvetica-Bold",spaceBefore=18,spaceAfter=8,leading=20),
        "h2":ParagraphStyle("h2",fontSize=12,textColor=C_TEAL,fontName="Helvetica-Bold",spaceBefore=12,spaceAfter=6,leading=16),
        "bd":ParagraphStyle("bd",fontSize=10,textColor=C_DARK,fontName="Helvetica",leading=15,spaceAfter=6,alignment=TA_JUSTIFY),
        "bk":ParagraphStyle("bk",fontSize=10,textColor=C_DARK,fontName="Helvetica-Bold",leading=14,spaceAfter=4),
    }
    def hr(): return HRFlowable(width="100%",thickness=1,color=C_TEAL,spaceAfter=10)
    def sp(h=8): return Spacer(1,h)
    def mtable(headers,rows,cw,hc=C_NAVY):
        data=[[Paragraph(f"<b>{h}</b>",ParagraphStyle("th",fontSize=9,textColor=colors.white,
            fontName="Helvetica-Bold",alignment=TA_CENTER)) for h in headers]]
        for row in rows:
            data.append([Paragraph(str(c),ParagraphStyle("td",fontSize=9,textColor=C_DARK,
                fontName="Helvetica",leading=13)) for c in row])
        t=Table(data,colWidths=cw,repeatRows=1)
        t.setStyle(TableStyle([
            ('BACKGROUND',(0,0),(-1,0),hc),
            ('ROWBACKGROUNDS',(0,1),(-1,-1),[colors.white,C_GREY]),
            ('GRID',(0,0),(-1,-1),.5,colors.HexColor("#cccccc")),
            ('VALIGN',(0,0),(-1,-1),'MIDDLE'),
            ('TOPPADDING',(0,0),(-1,-1),5),('BOTTOMPADDING',(0,0),(-1,-1),5),
            ('LEFTPADDING',(0,0),(-1,-1),6),('RIGHTPADDING',(0,0),(-1,-1),6),
        ]))
        return t
    now=datetime.now().strftime("%d %B %Y")
    story=[]
    # Cover
    story.append(sp(60))
    cd=[[Paragraph("DRUG STORAGE CONDITIONS SURVEY",sty["ct"])],
        [Paragraph("Research Report",sty["cs"])],[sp(10)],
        [Paragraph("Evaluation of Storage Conditions of Drugs in Community<br/>Pharmacies and Their Influence on Drug Potency and<br/>Shelf Life in the Sunyani Municipality, Ghana",sty["cs"])],
        [sp(20)],
        [Paragraph(f"Total Responses Analysed: <b>{n}</b>",sty["cs"])],
        [Paragraph(f"Report Generated: {now}",sty["cs"])],[sp(30)],
        [Paragraph("Sunyani Technical University | Department of Pharmacy",sty["cs"])],
        [Paragraph("Supervisor: Mrs. Lydia Sarfo Mainoo",sty["cs"])],
        [Paragraph("Researchers: Obeng Theophilus · Yussif Asmau · Egawu Naomi",sty["cs"])],]
    ct=Table(cd,colWidths=[17*cm])
    ct.setStyle(TableStyle([('BACKGROUND',(0,0),(-1,-1),C_NAVY),
        ('TOPPADDING',(0,0),(-1,-1),6),('BOTTOMPADDING',(0,0),(-1,-1),6),
        ('LEFTPADDING',(0,0),(-1,-1),20),('RIGHTPADDING',(0,0),(-1,-1),20)]))
    story+=[ct,PageBreak()]
    # Stats
    ac_p=df['has_ac'].mean()*100; deg_p=df['observed_degradation'].mean()*100
    gsp_p=df['gsp_training_received'].mean()*100; sop_p=df['has_written_sop'].mean()*100
    fda_p=df['fda_inspected'].mean()*100; nev_p=(df['temp_monitoring_freq']=="Never monitored").mean()*100
    # S1 Executive Summary
    story+=[Paragraph("1. Executive Summary",sty["h1"]),hr()]
    story.append(Paragraph(f"This report presents findings from a survey of <b>{n} community pharmacies</b> in "
        f"the Sunyani Municipality, Ghana, conducted in 2026 by the Department of Pharmacy, Sunyani Technical University. "
        f"The study evaluated drug storage conditions, WHO GSP compliance, and observed drug quality outcomes.",sty["bd"]))
    story+=[sp(),Paragraph("Key findings at a glance:",sty["bk"]),sp(4)]
    story.append(mtable(["Indicator","Value","Status"],[
        ["Total pharmacies surveyed",str(n),"—"],
        ["Have air conditioning",f"{ac_p:.1f}%","Below 70% benchmark" if ac_p<70 else "Above benchmark"],
        ["GSP training received",f"{gsp_p:.1f}%","Training gap" if gsp_p<70 else "Adequate"],
        ["Have written SOP",f"{sop_p:.1f}%","Procedural gap" if sop_p<70 else "Adequate"],
        ["Observed drug degradation",f"{deg_p:.1f}%","Active quality risk"],
        ["Never monitor temperature",f"{nev_p:.1f}%","Critical monitoring gap"],
        ["FDA-inspected recently",f"{fda_p:.1f}%","Regulatory oversight"],
    ],[8*cm,3*cm,6*cm]))
    story+=[sp(10),PageBreak()]
    # S2 Respondent Profile
    story+=[Paragraph("2. Respondent Profile",sty["h1"]),hr()]
    story.append(Paragraph("2.1 Pharmacy Type Distribution",sty["h2"]))
    pt=df["pharmacy_type"].value_counts().reset_index(); pt.columns=["Type","Count"]
    story.append(mtable(["Pharmacy Type","Count","Percentage"],
        [[r["Type"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in pt.iterrows()],[9*cm,3*cm,5*cm]))
    story+=[sp(6),Paragraph("2.2 Location Distribution",sty["h2"])]
    lt=df["location_type"].value_counts().reset_index(); lt.columns=["Location","Count"]
    story.append(mtable(["Location","Count","Percentage"],
        [[r["Location"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in lt.iterrows()],[9*cm,3*cm,5*cm]))
    story+=[sp(6),Paragraph("2.3 Years of Professional Experience",sty["h2"])]
    exp=df["years_experience"].describe().round(1)
    story.append(mtable(["Statistic","Value"],
        [["Mean",f"{exp['mean']} yrs"],["Median",f"{exp['50%']} yrs"],
         ["Min",f"{exp['min']} yrs"],["Max",f"{exp['max']} yrs"],
         ["Std Dev",f"{exp['std']} yrs"]],[9*cm,8*cm]))
    story+=[sp(10),PageBreak()]
    # S3 Infrastructure
    story+=[Paragraph("3. Storage Infrastructure Analysis",sty["h1"]),hr()]
    story.append(Paragraph("Table 3.1 shows availability of key storage equipment. The 70% benchmark represents minimum adequate coverage.",sty["bd"]))
    story.append(sp(6))
    ek=["has_ac","has_refrigerator","has_thermometer","has_hygrometer","has_proper_shelving","has_ventilation"]
    el=["Air Conditioning","Refrigerator","Thermometer","Hygrometer","Proper Shelving","Ventilation"]
    story.append(mtable(["Equipment","Count","% of Pharmacies","Status"],
        [[lbl,str(int(df[k].sum())),f"{df[k].mean()*100:.1f}%","Adequate" if df[k].mean()*100>=70 else "Below Benchmark"]
         for k,lbl in zip(ek,el)],[6*cm,2.5*cm,3.5*cm,5*cm]))
    story.append(Paragraph(f"<b>Interpretation:</b> Air conditioning is available in {ac_p:.1f}% of pharmacies surveyed — "
        f"{'above' if ac_p>=70 else 'below'} the 70% benchmark. This directly affects the ability to maintain "
        f"WHO-recommended storage temperatures (25-30°C).",sty["bd"]))
    story+=[sp(10),PageBreak()]
    # S4 Temperature
    story+=[Paragraph("4. Temperature and Humidity Monitoring",sty["h1"]),hr()]
    story.append(Paragraph("4.1 Temperature Range Distribution",sty["h2"]))
    tm2={"Below 25°C  (WHO Optimal)":"Below 25C (WHO Optimal)","25°C – 30°C  (Acceptable)":"25-30C (Acceptable)",
         "30°C – 35°C  (Risk Zone)":"30-35C (Risk Zone)","Above 35°C  (High Risk)":"Above 35C (High Risk)","Not monitored / Unknown":"Not Monitored"}
    df["tl2"]=df["usual_temp_range"].map(tm2).fillna(df["usual_temp_range"])
    tc2=df["tl2"].value_counts().reset_index(); tc2.columns=["Range","Count"]
    story.append(mtable(["Temperature Range","Count","% of Pharmacies"],
        [[r["Range"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in tc2.iterrows()],[8*cm,3*cm,6*cm]))
    trisk=df[df["usual_temp_range"].isin(["30°C – 35°C  (Risk Zone)","Above 35°C  (High Risk)"])]
    nmon=df[df["temp_monitoring_freq"]=="Never monitored"]
    story.append(Paragraph(f"<b>Finding:</b> {len(trisk)} pharmacies ({len(trisk)/n*100:.1f}%) operate above WHO-recommended temperature limits. "
        f"Additionally, {len(nmon)} ({len(nmon)/n*100:.1f}%) never monitor temperature at all.",sty["bd"]))
    story+=[sp(8),Paragraph("4.2 Humidity Range Distribution",sty["h2"])]
    hm2={"Below 45% RH  (Low)":"Below 45% RH (Low)","45% – 65% RH  (Optimal)":"45-65% RH (Optimal)",
         "65% – 75% RH  (Elevated Risk)":"65-75% RH (Elevated Risk)","Above 75% RH  (High Risk)":"Above 75% RH (High Risk)","Not monitored / Unknown":"Not Monitored"}
    df["hl2"]=df["usual_humidity_range"].map(hm2).fillna(df["usual_humidity_range"])
    hc2=df["hl2"].value_counts().reset_index(); hc2.columns=["Humidity","Count"]
    story.append(mtable(["Humidity Range","Count","% of Pharmacies"],
        [[r["Humidity"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in hc2.iterrows()],[8*cm,3*cm,6*cm]))
    hrisk=df[df["usual_humidity_range"].isin(["65% – 75% RH  (Elevated Risk)","Above 75% RH  (High Risk)"])]
    story.append(Paragraph(f"<b>Finding:</b> {len(hrisk)} pharmacies ({len(hrisk)/n*100:.1f}%) reported humidity above 65% RH, "
        f"accelerating moisture-induced degradation in oral tablets, capsules and antimalarials.",sty["bd"]))
    story+=[sp(10),PageBreak()]
    # S5 GSP
    story+=[Paragraph("5. WHO Good Storage Practice (GSP) Compliance",sty["h1"]),hr()]
    story.append(Paragraph("A compliance rate below 70% is classified as non-compliant at the population level.",sty["bd"]))
    story.append(sp(6))
    gsp_ind=[
        ("GSP Training Received",df["gsp_training_received"].mean()*100,int(df["gsp_training_received"].sum())),
        ("Written SOP Exists",df["has_written_sop"].mean()*100,int(df["has_written_sop"].sum())),
        ("FIFO Practised (Always)",(df["practises_fifo"]=="Yes").mean()*100,int((df["practises_fifo"]=="Yes").sum())),
        ("Expiry Checks (Always)",(df["checks_expiry"]=="Yes").mean()*100,int((df["checks_expiry"]=="Yes").sum())),
        ("Cold Chain for Thermolabiles",(df["segregates_thermolabile"]=="Yes").mean()*100,int((df["segregates_thermolabile"]=="Yes").sum())),
        ("FDA-Inspected (Last 2 Yrs)",df["fda_inspected"].mean()*100,int(df["fda_inspected"].sum())),
    ]
    story.append(mtable(["GSP Indicator","Compliant (n)","Non-compliant (n)","% Compliant","Level"],
        [[lbl,str(cnt),str(n-cnt),f"{pct:.1f}%","Good" if pct>=70 else ("Moderate" if pct>=40 else "Poor")]
         for lbl,pct,cnt in gsp_ind],[6.5*cm,2.5*cm,2.8*cm,2.5*cm,2.7*cm]))
    story+=[sp(10),PageBreak()]
    # S6 Quality
    story+=[Paragraph("6. Observed Drug Quality and Potency Issues",sty["h1"]),hr()]
    story.append(mtable(["Quality Indicator","Yes (n)","No (n)","% Yes","Research Implication"],
        [["Physical Drug Degradation Observed",str(int(df["observed_degradation"].sum())),str(n-int(df["observed_degradation"].sum())),
          f"{df['observed_degradation'].mean()*100:.1f}%","Active drug deterioration in storage"],
         ["Unusual Colour Change Noted",str(int(df["observed_color_change"].sum())),str(n-int(df["observed_color_change"].sum())),
          f"{df['observed_color_change'].mean()*100:.1f}%","Chemical/physical instability"],
         ["Patient Potency Complaints",str(int(df["potency_complaints"].sum())),str(n-int(df["potency_complaints"].sum())),
          f"{df['potency_complaints'].mean()*100:.1f}%","Possible therapeutic failure"],
         ["Stock Returned/Discarded",str(int(df["returned_stock_quality"].sum())),str(n-int(df["returned_stock_quality"].sum())),
          f"{df['returned_stock_quality'].mean()*100:.1f}%","Economic loss indicator"],
        ],[5.5*cm,1.8*cm,1.8*cm,1.8*cm,6.1*cm]))
    a2=[]
    for row in df["degradation_drug_types"].dropna():
        a2.extend([t.strip() for t in str(row).split(",") if t.strip()])
    if a2:
        story+=[sp(8),Paragraph("6.1 Drug Categories Affected by Degradation",sty["h2"])]
        tc_d=Counter(a2)
        dr=sorted([[k,str(v),f"{v/n*100:.1f}%"] for k,v in tc_d.items()],key=lambda x:-int(x[1]))
        story.append(mtable(["Drug Category","Reports","% of Pharmacies"],dr,[8*cm,3*cm,6*cm]))
    story+=[sp(10),PageBreak()]
    # S7 Key Findings
    story+=[Paragraph("7. Summary of Key Research Findings",sty["h1"]),hr()]
    finds=[
        f"Only {ac_p:.1f}% of pharmacies have air conditioning — {'above' if ac_p>=70 else 'below'} the 70% benchmark.",
        f"{gsp_p:.1f}% of pharmacists have received WHO GSP training, indicating a significant knowledge gap.",
        f"Only {sop_p:.1f}% of pharmacies have a written Standard Operating Procedure for drug storage.",
        f"{deg_p:.1f}% of pharmacies have observed physical signs of drug degradation.",
        f"{df['potency_complaints'].mean()*100:.1f}% have received patient complaints about drug ineffectiveness.",
        f"{nev_p:.1f}% of pharmacies never monitor storage temperature — the most critical monitoring failure.",
        f"{fda_p:.1f}% have been inspected by FDA Ghana in the last 2 years.",
        f"Top challenge: '{df['biggest_challenge'].mode()[0] if not df['biggest_challenge'].empty else 'N/A'}'.",
        f"Top requested support: '{df['support_needed'].mode()[0] if not df['support_needed'].empty else 'N/A'}'.",
    ]
    for i,f_ in enumerate(finds,1):
        story.append(Paragraph(f"{i}. {f_}",sty["bd"])); story.append(sp(3))
    story+=[sp(10),PageBreak()]
    # S8 Recommendations
    story+=[Paragraph("8. Evidence-Based Policy Recommendations",sty["h1"]),hr()]
    recs=[
        ("Mandatory Temperature Monitoring Equipment",
         "The FDA Ghana should require all licensed pharmacies to own calibrated thermometers and hygrometers "
         "as a condition for annual operating licence renewal."),
        ("Subsidised Air Conditioning Access",
         f"Government and pharmaceutical associations should co-fund AC installation — particularly for the "
         f"{100-ac_p:.0f}% of pharmacies currently without it."),
        ("Compulsory Annual WHO GSP Training",
         "The Pharmacy Council of Ghana should mandate annual GSP refresher training for all registered pharmacists, "
         "with digital certification upon completion."),
        ("Standardised SOP Templates",
         "The Pharmacy Council should distribute standardised, ready-to-use drug storage SOPs freely to all community pharmacies."),
        ("Risk-Based FDA Inspection Increase",
         "A risk-stratified inspection schedule should prioritise rural and peri-urban pharmacies with known temperature or humidity challenges."),
        ("Community Pharmacy Infrastructure Grants",
         "The Ministry of Health should establish a dedicated grant programme for facility upgrades in resource-limited pharmacies."),
    ]
    for i,(title,body) in enumerate(recs,1):
        story.append(Paragraph(f"{i}. {title}",sty["h2"]))
        story.append(Paragraph(body,sty["bd"])); story.append(sp(4))
    # S9 Verbatim
    pr=df["policy_recommendation"].dropna()
    pr=pr[pr.str.strip()!=""]
    if not pr.empty:
        story+=[sp(8),PageBreak(),Paragraph("9. Verbatim Policy Recommendations from Pharmacists",sty["h1"]),hr()]
        story.append(Paragraph("The following were provided directly by survey respondents:",sty["bd"])); story.append(sp(6))
        for i,rec in enumerate(pr[:20],1):
            story.append(Paragraph(f"{i}. {rec}",sty["bd"])); story.append(sp(3))
    # Footer
    story+=[sp(20),hr(),
        Paragraph("Report automatically generated by the Drug Storage Conditions Survey Platform",
            ParagraphStyle("ft",fontSize=9,textColor=colors.grey,alignment=TA_CENTER)),
        Paragraph(f"Sunyani Technical University — Department of Pharmacy — {now}",
            ParagraphStyle("ft2",fontSize=9,textColor=colors.grey,alignment=TA_CENTER))]
    doc.build(story)
    buf.seek(0)
    return buf.read()

# ════════════════════════════════════════════════════════
# WORD GENERATOR
# ════════════════════════════════════════════════════════
def generate_word(df,n):
    from docx import Document
    from docx.shared import Pt,Cm,RGBColor
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    doc=Document()
    for sec in doc.sections:
        sec.top_margin=Cm(2.5); sec.bottom_margin=Cm(2.0)
        sec.left_margin=Cm(2.5); sec.right_margin=Cm(2.5)
    def rgb(h):
        h=h.lstrip("#"); return RGBColor(int(h[:2],16),int(h[2:4],16),int(h[4:],16))
    def set_bg(cell,hx):
        tc=cell._tc; tcPr=tc.get_or_add_tcPr()
        shd=OxmlElement('w:shd'); shd.set(qn('w:val'),'clear')
        shd.set(qn('w:color'),'auto'); shd.set(qn('w:fill'),hx.lstrip('#'))
        tcPr.append(shd)
    def head(text,level=1,color="065A82"):
        p=doc.add_heading(text,level=level)
        run=p.runs[0] if p.runs else p.add_run(text)
        run.font.color.rgb=rgb(color); run.font.name="Calibri"
        p.paragraph_format.space_before=Pt(14); p.paragraph_format.space_after=Pt(6)
        return p
    def body(text,bold=False,color=None,italic=False):
        p=doc.add_paragraph(); run=p.add_run(text)
        run.font.size=Pt(10.5); run.font.name="Calibri"
        run.bold=bold; run.italic=italic
        if color: run.font.color.rgb=rgb(color)
        p.paragraph_format.space_after=Pt(4); p.paragraph_format.space_before=Pt(0)
        return p
    def add_table(headers,rows,cw,hc="065A82"):
        t=doc.add_table(rows=1+len(rows),cols=len(headers)); t.style="Table Grid"
        hr_=t.rows[0]
        for i,h in enumerate(headers):
            cell=hr_.cells[i]; cell.width=Cm(cw[i]); set_bg(cell,hc)
            p=cell.paragraphs[0]; run=p.add_run(h)
            run.bold=True; run.font.color.rgb=rgb("FFFFFF")
            run.font.size=Pt(9.5); run.font.name="Calibri"
            p.alignment=WD_ALIGN_PARAGRAPH.CENTER; cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        for ri,row in enumerate(rows):
            tr=t.rows[ri+1]; bg="F5F5F5" if ri%2==0 else "FFFFFF"
            for ci,cv in enumerate(row):
                cell=tr.cells[ci]; cell.width=Cm(cw[ci]); set_bg(cell,bg)
                p=cell.paragraphs[0]; run=p.add_run(str(cv))
                run.font.size=Pt(9.5); run.font.name="Calibri"
                cell.vertical_alignment=WD_ALIGN_VERTICAL.CENTER
        doc.add_paragraph(); return t
    def add_hr():
        p=doc.add_paragraph(); pPr=p._p.get_or_add_pPr()
        pBdr=OxmlElement('w:pBdr'); bot=OxmlElement('w:bottom')
        bot.set(qn('w:val'),'single'); bot.set(qn('w:sz'),'6'); bot.set(qn('w:color'),'028090')
        pBdr.append(bot); pPr.append(pBdr); p.paragraph_format.space_after=Pt(6)
    now=datetime.now().strftime("%d %B %Y")
    ac_p=df['has_ac'].mean()*100; deg_p=df['observed_degradation'].mean()*100
    gsp_p=df['gsp_training_received'].mean()*100; sop_p=df['has_written_sop'].mean()*100
    fda_p=df['fda_inspected'].mean()*100; nev_p=(df['temp_monitoring_freq']=="Never monitored").mean()*100
    # Cover
    doc.add_paragraph()
    tp=doc.add_paragraph(); tp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run=tp.add_run("DRUG STORAGE CONDITIONS SURVEY"); run.bold=True
    run.font.size=Pt(22); run.font.name="Calibri"; run.font.color.rgb=rgb("028090")
    sp2=doc.add_paragraph(); sp2.alignment=WD_ALIGN_PARAGRAPH.CENTER
    run2=sp2.add_run("Research Report — Department of Pharmacy")
    run2.font.size=Pt(13); run2.font.name="Calibri"; run2.font.color.rgb=rgb("065A82")
    doc.add_paragraph()
    for line in ["Evaluation of Storage Conditions of Drugs in Community Pharmacies",
        "and Their Influence on Drug Potency and Shelf Life","in the Sunyani Municipality, Ghana","",
        f"Total Responses: {n}   |   Report Date: {now}","",
        "Sunyani Technical University | Department of Pharmacy",
        "Supervisor: Mrs. Lydia Sarfo Mainoo","Researchers: Obeng Theophilus · Yussif Asmau · Egawu Naomi"]:
        p=doc.add_paragraph(line); p.alignment=WD_ALIGN_PARAGRAPH.CENTER
        p.paragraph_format.space_after=Pt(2)
        if p.runs: p.runs[0].font.name="Calibri"
    doc.add_page_break()
    # S1
    head("1. Executive Summary"); add_hr()
    body(f"This report presents findings from a survey of {n} community pharmacies in the Sunyani Municipality, Ghana, conducted in 2026 by the Department of Pharmacy, Sunyani Technical University.")
    doc.add_paragraph()
    add_table(["Indicator","Value","Status"],
        [["Total pharmacies surveyed",str(n),"—"],
         ["Have air conditioning",f"{ac_p:.1f}%","Below 70% benchmark" if ac_p<70 else "Above benchmark"],
         ["GSP training received",f"{gsp_p:.1f}%","Training gap" if gsp_p<70 else "Adequate"],
         ["Have written SOP",f"{sop_p:.1f}%","Procedural gap" if sop_p<70 else "Adequate"],
         ["Observed drug degradation",f"{deg_p:.1f}%","Active quality risk"],
         ["Never monitor temperature",f"{nev_p:.1f}%","Critical monitoring gap"],
         ["FDA-inspected recently",f"{fda_p:.1f}%","Regulatory oversight"]],[9,3.5,5])
    doc.add_page_break()
    # S2
    head("2. Respondent Profile"); add_hr()
    head("2.1 Pharmacy Type Distribution",level=2,color="028090")
    pt=df["pharmacy_type"].value_counts().reset_index(); pt.columns=["Type","Count"]
    add_table(["Pharmacy Type","Count","Percentage"],
        [[r["Type"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in pt.iterrows()],[9,3,5.5])
    head("2.2 Location Distribution",level=2,color="028090")
    lt=df["location_type"].value_counts().reset_index(); lt.columns=["Location","Count"]
    add_table(["Location","Count","Percentage"],
        [[r["Location"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in lt.iterrows()],[9,3,5.5])
    head("2.3 Professional Experience",level=2,color="028090")
    exp=df["years_experience"].describe().round(1)
    add_table(["Statistic","Value"],[["Mean",f"{exp['mean']} years"],["Median",f"{exp['50%']} years"],
        ["Min",f"{exp['min']} years"],["Max",f"{exp['max']} years"],["Std Dev",f"{exp['std']} years"]],[9,8.5])
    doc.add_page_break()
    # S3
    head("3. Storage Infrastructure Analysis"); add_hr()
    body("Table 3.1 shows availability of essential storage equipment. The 70% benchmark reflects minimum adequate coverage.")
    doc.add_paragraph()
    ek=["has_ac","has_refrigerator","has_thermometer","has_hygrometer","has_proper_shelving","has_ventilation"]
    el=["Air Conditioning","Refrigerator","Thermometer","Hygrometer","Proper Shelving","Ventilation"]
    add_table(["Equipment","Count","% of Pharmacies","Status"],
        [[lbl,str(int(df[k].sum())),f"{df[k].mean()*100:.1f}%","Adequate" if df[k].mean()*100>=70 else "Below Benchmark"]
         for k,lbl in zip(ek,el)],[5.5,2.5,3.5,6])
    body(f"Air conditioning is present in {ac_p:.1f}% of pharmacies — {'above' if ac_p>=70 else 'below'} the 70% benchmark.",bold=False)
    doc.add_page_break()
    # S4
    head("4. Temperature and Humidity Monitoring"); add_hr()
    head("4.1 Temperature Ranges",level=2,color="028090")
    tm2={"Below 25°C  (WHO Optimal)":"Below 25C (WHO Optimal)","25°C – 30°C  (Acceptable)":"25-30C (Acceptable)",
         "30°C – 35°C  (Risk Zone)":"30-35C (Risk Zone)","Above 35°C  (High Risk)":"Above 35C (High Risk)","Not monitored / Unknown":"Not Monitored"}
    df["tl2"]=df["usual_temp_range"].map(tm2).fillna(df["usual_temp_range"])
    tc2=df["tl2"].value_counts().reset_index(); tc2.columns=["Range","Count"]
    add_table(["Temperature Range","Count","% of Pharmacies"],
        [[r["Range"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in tc2.iterrows()],[8,3,6.5])
    trisk=df[df["usual_temp_range"].isin(["30°C – 35°C  (Risk Zone)","Above 35°C  (High Risk)"])]
    nmon=df[df["temp_monitoring_freq"]=="Never monitored"]
    body(f"Finding: {len(trisk)} pharmacies ({len(trisk)/n*100:.1f}%) operate above WHO safe temperature limits. {len(nmon)} ({len(nmon)/n*100:.1f}%) never monitor temperature.")
    head("4.2 Humidity Ranges",level=2,color="028090")
    hm2={"Below 45% RH  (Low)":"Below 45% RH","45% – 65% RH  (Optimal)":"45-65% RH (Optimal)",
         "65% – 75% RH  (Elevated Risk)":"65-75% RH (Elevated Risk)","Above 75% RH  (High Risk)":"Above 75% RH (High Risk)","Not monitored / Unknown":"Not Monitored"}
    df["hl2"]=df["usual_humidity_range"].map(hm2).fillna(df["usual_humidity_range"])
    hc2=df["hl2"].value_counts().reset_index(); hc2.columns=["Humidity","Count"]
    add_table(["Humidity Range","Count","% of Pharmacies"],
        [[r["Humidity"],str(r["Count"]),f"{r['Count']/n*100:.1f}%"] for _,r in hc2.iterrows()],[8,3,6.5])
    doc.add_page_break()
    # S5
    head("5. WHO GSP Compliance"); add_hr()
    gsp_ind=[
        ("GSP Training Received",df["gsp_training_received"].mean()*100,int(df["gsp_training_received"].sum())),
        ("Written SOP Exists",df["has_written_sop"].mean()*100,int(df["has_written_sop"].sum())),
        ("FIFO Practised (Always)",(df["practises_fifo"]=="Yes").mean()*100,int((df["practises_fifo"]=="Yes").sum())),
        ("Expiry Checks (Always)",(df["checks_expiry"]=="Yes").mean()*100,int((df["checks_expiry"]=="Yes").sum())),
        ("Cold Chain for Thermolabiles",(df["segregates_thermolabile"]=="Yes").mean()*100,int((df["segregates_thermolabile"]=="Yes").sum())),
        ("FDA-Inspected (2 Years)",df["fda_inspected"].mean()*100,int(df["fda_inspected"].sum())),
    ]
    add_table(["GSP Indicator","Compliant (n)","Non-compliant (n)","% Compliant","Level"],
        [[lbl,str(cnt),str(n-cnt),f"{pct:.1f}%","Good" if pct>=70 else ("Moderate" if pct>=40 else "Poor")]
         for lbl,pct,cnt in gsp_ind],[6,2.5,2.8,2.5,3.7])
    doc.add_page_break()
    # S6
    head("6. Observed Drug Quality and Potency Issues"); add_hr()
    add_table(["Indicator","Yes (n)","% Yes","Implication"],
        [["Physical Degradation Observed",str(int(df["observed_degradation"].sum())),f"{df['observed_degradation'].mean()*100:.1f}%","Active deterioration in storage"],
         ["Unusual Colour Change",str(int(df["observed_color_change"].sum())),f"{df['observed_color_change'].mean()*100:.1f}%","Chemical/physical instability"],
         ["Patient Potency Complaints",str(int(df["potency_complaints"].sum())),f"{df['potency_complaints'].mean()*100:.1f}%","Possible therapeutic failure"],
         ["Stock Returned/Discarded",str(int(df["returned_stock_quality"].sum())),f"{df['returned_stock_quality'].mean()*100:.1f}%","Economic loss indicator"]],[6,2,2.5,7])
    a2=[]
    for row in df["degradation_drug_types"].dropna():
        a2.extend([t.strip() for t in str(row).split(",") if t.strip()])
    if a2:
        head("6.1 Drug Categories Affected",level=2,color="028090")
        tc_d=Counter(a2)
        dr=sorted([[k,str(v),f"{v/n*100:.1f}%"] for k,v in tc_d.items()],key=lambda x:-int(x[1]))
        add_table(["Drug Category","Reports","% of Pharmacies"],dr,[8,3,6.5])
    doc.add_page_break()
    # S7
    head("7. Key Research Findings"); add_hr()
    finds=[f"Only {ac_p:.1f}% of pharmacies have air conditioning.",
        f"{gsp_p:.1f}% of pharmacists have received WHO GSP training.",
        f"Only {sop_p:.1f}% of pharmacies have a written drug storage SOP.",
        f"{deg_p:.1f}% of pharmacies have observed physical drug degradation.",
        f"{df['potency_complaints'].mean()*100:.1f}% have received patient potency complaints.",
        f"{nev_p:.1f}% of pharmacies never monitor storage temperature.",
        f"{fda_p:.1f}% have been FDA-inspected in the last 2 years."]
    for i,f_ in enumerate(finds,1):
        p=doc.add_paragraph(style="List Number"); run=p.add_run(f_)
        run.font.size=Pt(10.5); run.font.name="Calibri"
    doc.add_page_break()
    # S8
    head("8. Evidence-Based Policy Recommendations"); add_hr()
    recs=[
        ("Mandatory Temperature Monitoring Equipment","The FDA Ghana should require calibrated thermometers and hygrometers as a condition of annual licence renewal."),
        ("Subsidised Air Conditioning Access",f"Government should co-fund AC installation, prioritising the {100-ac_p:.0f}% of pharmacies without it."),
        ("Compulsory Annual WHO GSP Training","Mandate annual refresher training for all registered pharmacists through the Pharmacy Council."),
        ("Standardised SOP Templates","Distribute ready-to-use drug storage SOPs freely to all community pharmacies."),
        ("Risk-Based FDA Inspection Increase","Prioritise rural and peri-urban pharmacies with known temperature or humidity challenges."),
        ("Community Pharmacy Infrastructure Grants","Ministry of Health should create a grant programme for facility upgrades in resource-limited pharmacies."),
    ]
    for i,(title,rec_body) in enumerate(recs,1):
        head(f"{i}. {title}",level=2,color="028090"); body(rec_body)
    pr=df["policy_recommendation"].dropna(); pr=pr[pr.str.strip()!=""]
    if not pr.empty:
        doc.add_page_break(); head("9. Verbatim Recommendations from Pharmacists"); add_hr()
        body("The following were provided directly by survey respondents:"); doc.add_paragraph()
        for i,rec in enumerate(pr[:20],1):
            p=doc.add_paragraph(style="List Number"); run=p.add_run(rec)
            run.font.size=Pt(10.5); run.font.name="Calibri"; run.font.color.rgb=rgb("065A82")
    doc.add_paragraph(); add_hr()
    fp=doc.add_paragraph(f"Report generated · Sunyani Technical University · Department of Pharmacy · {now}")
    fp.alignment=WD_ALIGN_PARAGRAPH.CENTER
    if fp.runs: fp.runs[0].font.size=Pt(8.5); fp.runs[0].font.color.rgb=rgb("888888"); fp.runs[0].font.name="Calibri"
    buf=BytesIO(); doc.save(buf); buf.seek(0); return buf.read()


# ════════════════════════════════════════════════════════
# PAGE: REPORT (Admin)
# ════════════════════════════════════════════════════════
def page_report():
    if not require_admin(): return
    df=fetch_all()
    if df.empty: st.warning("📭 No data available. Submit some responses first."); return
    df=prep_df(df); n=len(df)
    now=datetime.now().strftime("%d %B %Y, %H:%M")
    st.markdown(f"## 📄 Thesis Research Report\n**Generated:** {now} &nbsp;|&nbsp; **Total Responses:** {n}\n---")
    ac_p=df['has_ac'].mean()*100; deg_p=df['observed_degradation'].mean()*100; gsp_p=df['gsp_training_received'].mean()*100
    sop_p=df['has_written_sop'].mean()*100; fda_p=df['fda_inspected'].mean()*100
    nev_p=(df['temp_monitoring_freq']=="Never monitored").mean()*100
    c1,c2,c3=st.columns(3)
    c1.metric("Air Conditioning Coverage",f"{ac_p:.1f}%","below benchmark" if ac_p<70 else "above benchmark")
    c2.metric("GSP Training Coverage",f"{gsp_p:.1f}%")
    c3.metric("Drug Degradation Observed",f"{deg_p:.1f}%")
    st.markdown("### 📥 Download Reports")
    col1,col2,col3=st.columns(3)
    with col1:
        with st.spinner("Generating PDF..."):
            try:
                pb=generate_pdf(df.copy(),n)
                st.download_button("📄 Download PDF Report",data=pb,
                    file_name=f"STU_DrugStorage_{datetime.now().strftime('%Y%m%d')}.pdf",
                    mime="application/pdf",use_container_width=True)
            except Exception as e: st.error(f"PDF error: {e}")
    with col2:
        with st.spinner("Generating Word doc..."):
            try:
                wb=generate_word(df.copy(),n)
                st.download_button("📝 Download Word Report",data=wb,
                    file_name=f"STU_DrugStorage_{datetime.now().strftime('%Y%m%d')}.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document",
                    use_container_width=True)
            except Exception as e: st.error(f"Word error: {e}")
    with col3:
        csv_d=df.to_csv(index=False).encode("utf-8")
        st.download_button("📊 Download CSV Data",data=csv_d,
            file_name=f"STU_DrugStorage_Data_{datetime.now().strftime('%Y%m%d')}.csv",
            mime="text/csv",use_container_width=True)
    st.markdown("---")
    st.markdown("### Report Preview")
    st.markdown("#### 1. Respondent Profile")
    c1,c2,c3=st.columns(3)
    with c1:
        st.markdown("**Pharmacy Type**")
        pt=df["pharmacy_type"].value_counts().reset_index(); pt.columns=["Type","Count"]
        pt["Percentage"]=(pt["Count"]/n*100).round(1).astype(str)+"%"
        st.dataframe(pt,hide_index=True,use_container_width=True)
    with c2:
        st.markdown("**Location**")
        lt=df["location_type"].value_counts().reset_index(); lt.columns=["Location","Count"]
        lt["Percentage"]=(lt["Count"]/n*100).round(1).astype(str)+"%"
        st.dataframe(lt,hide_index=True,use_container_width=True)
    with c3:
        st.markdown("**Experience (years)**")
        es=df["years_experience"].describe().round(1)
        st.dataframe(es.reset_index().rename(columns={"index":"Stat","years_experience":"Years"}),hide_index=True,use_container_width=True)
    st.markdown("#### 2. Storage Infrastructure")
    ek=["has_ac","has_refrigerator","has_thermometer","has_hygrometer","has_proper_shelving","has_ventilation"]
    el=["Air Conditioning","Refrigerator","Thermometer","Hygrometer","Proper Shelving","Ventilation"]
    st.dataframe(pd.DataFrame({"Equipment":el,"Count":[int(df[k].sum()) for k in ek],
        "% of Pharmacies":[f"{df[k].mean()*100:.1f}%" for k in ek],
        "Status":["✅ Adequate" if df[k].mean()*100>=70 else "❌ Below Benchmark" for k in ek]}),
        hide_index=True,use_container_width=True)
    st.markdown("#### 3. WHO GSP Compliance")
    gd={"GSP Indicator":["GSP Training","Written SOP","FIFO (Always)","Expiry Checks (Always)","Cold Chain","FDA Inspected"],
        "Compliant (n)":[int(df["gsp_training_received"].sum()),int(df["has_written_sop"].sum()),
            int((df["practises_fifo"]=="Yes").sum()),int((df["checks_expiry"]=="Yes").sum()),
            int((df["segregates_thermolabile"]=="Yes").sum()),int(df["fda_inspected"].sum())]}
    gdf2=pd.DataFrame(gd)
    gdf2["% Compliant"]=(gdf2["Compliant (n)"]/n*100).round(1).astype(str)+"%"
    gdf2["Non-compliant (n)"]=n-gdf2["Compliant (n)"]
    gdf2["Level"]=gdf2["Compliant (n)"].apply(lambda x:"✅ Good" if x/n>=.7 else ("⚠️ Moderate" if x/n>=.4 else "❌ Poor"))
    st.dataframe(gdf2,hide_index=True,use_container_width=True)
    st.markdown("#### 4. Key Findings")
    for i,f_ in enumerate([f"Air conditioning in {ac_p:.1f}% — {'above' if ac_p>=70 else 'below'} 70% benchmark.",
        f"{gsp_p:.1f}% have WHO GSP training.",f"Only {sop_p:.1f}% have a written SOP.",
        f"{deg_p:.1f}% observed drug degradation.",f"{df['potency_complaints'].mean()*100:.1f}% received potency complaints.",
        f"{nev_p:.1f}% never monitor temperature.",f"{fda_p:.1f}% FDA-inspected in 2 years."],1):
        st.markdown(f"**{i}.** {f_}")
    st.markdown("#### 5. Policy Recommendations")
    st.markdown("""
1. **Mandatory monitoring equipment** — Require thermometers and hygrometers for licence renewal.
2. **Subsidise air conditioning** — Government co-funding for AC in under-resourced pharmacies.
3. **Annual GSP training** — Mandate annual refresher training through the Pharmacy Council.
4. **Free SOP templates** — Distribute standardised storage SOPs to all community pharmacies.
5. **Risk-based FDA inspections** — Increase frequency with priority on rural and high-risk facilities.
6. **Infrastructure grants** — Ministry of Health grant programme for facility upgrades.
    """)
    pr=df["policy_recommendation"].dropna(); pr=pr[pr.str.strip()!=""]
    if not pr.empty:
        st.markdown("#### 6. Verbatim Recommendations from Pharmacists")
        for i,rec in enumerate(pr[:20],1): st.markdown(f"**{i}.** _{rec}_")


# ════════════════════════════════════════════════════════
# PAGE: ABOUT
# ════════════════════════════════════════════════════════
def page_about():
    st.markdown("""
    ## 📖 About This Survey System
    > **"Evaluation of Storage Conditions of Drugs in Community Pharmacies and
    > Their Influence on Drug Potency and Shelf Life in the Sunyani Municipality, Ghana"**
    ---
    ### 🎓 Academic Details
    | | |
    |---|---|
    | **Institution** | Sunyani Technical University |
    | **Department** | Department of Pharmacy |
    | **Supervisor** | Mrs. Lydia Sarfo Mainoo |
    | **Researchers** | Obeng Theophilus (STUBTECH220600) · Yussif Asmau (STUBTECH220592) · Egawu Naomi (STUBTECH220598) |
    | **Year** | 2026 |
    ---
    ### 🗄️ Tech Stack
    | Component | Technology |
    |-----------|------------|
    | UI | Streamlit (Streamlit Cloud) |
    | Database | Supabase (PostgreSQL) |
    | Charts | Plotly (interactive) |
    | PDF Report | ReportLab |
    | Word Report | python-docx |
    | Language | Python 3.x |
    ---
    ### 🔒 Data Privacy
    All responses are stored securely in Supabase PostgreSQL.
    Data is used exclusively for academic research.
    The Admin Panel is password-protected.
    """)


# ════════════════════════════════════════════════════════
# NAVIGATION
# ════════════════════════════════════════════════════════
render_header()

st.sidebar.markdown("""
<div style="text-align:center;padding:10px 0 18px;">
    <div style="font-size:2.4em;">💊</div>
    <div style="font-weight:800;color:#028090;font-size:1.0em;">Drug Storage Survey</div>
    <div style="font-size:.75em;color:#888;margin-top:4px;">
        Sunyani Technical University<br>Department of Pharmacy
    </div>
</div>
""",unsafe_allow_html=True)

page=st.sidebar.radio("Navigation",
    ["📋  Fill Survey","📊  Dashboard & Charts","📁  Responses (Admin)","📄  Report & Downloads (Admin)","ℹ️  About"],
    label_visibility="collapsed")

st.sidebar.markdown("---")
count=response_count()
st.sidebar.markdown(f"""
<div style="background:linear-gradient(135deg,#028090,#065A82);
            color:white;border-radius:10px;padding:14px;text-align:center;">
    <div style="font-size:2em;font-weight:800;">{count}</div>
    <div style="font-size:.82em;opacity:.9;">Responses Collected</div>
</div>
""",unsafe_allow_html=True)

st.sidebar.markdown("""
<div style="margin-top:18px;font-size:.76em;color:#999;text-align:center;line-height:1.6;">
    <b>Database:</b> Supabase<br>
    <b>Hosting:</b> Streamlit Cloud<br><br>
    <b>Admin panel</b> is password protected.
</div>
""",unsafe_allow_html=True)

if   "Fill Survey" in page: page_survey()
elif "Dashboard"   in page: page_dashboard()
elif "Responses"   in page: page_responses()
elif "Report"      in page: page_report()
elif "About"       in page: page_about()
